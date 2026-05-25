"""
generar_doc.py — Genera documentación completa MODULO 5 AGP Glass
Ejecutar: py generar_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Estilos base ──────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def set_heading(paragraph, level=1):
    colors = {1: "1F3864", 2: "2E5299", 3: "4472C4", 4: "5B9BD5"}
    sizes  = {1: 20, 2: 16, 3: 13, 4: 12}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    run.font.color.rgb = RGBColor.from_string(colors.get(level, "000000"))
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True

def h1(text):
    p = doc.add_heading(text, level=1)
    set_heading(p, 1)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    set_heading(p, 2)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    set_heading(p, 3)
    return p

def h4(text):
    p = doc.add_heading(text, level=4)
    set_heading(p, 4)
    return p

def para(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("2F5496")
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(level * 0.5 + 0.5)
    return p

def tabla(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor.from_string("FFFFFF")
        # Fondo azul oscuro
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3864')
        tcPr.append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Datos
    for ri, row_data in enumerate(rows):
        row_obj = t.rows[ri + 1]
        fill = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row_obj.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), fill)
            tcPr.append(shd)
    if col_widths:
        n_cols = len(headers)
        for ci, w in enumerate(col_widths):
            if ci >= n_cols:
                break
            for row_obj in t.rows:
                row_obj.cells[ci].width = Cm(w)
    return t

def separador():
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════
p_titulo = doc.add_paragraph()
p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_titulo.add_run("MÓDULO 5 — AGP GLASS COLOMBIA")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor.from_string("1F3864")

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run("Sistema de Automatización SAP y Gestión de ZFERs")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor.from_string("2E5299")
run2.italic = True

separador()
p_fecha = doc.add_paragraph()
p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fecha.add_run(f"Documentación técnica generada el {datetime.date.today().strftime('%d de %B de %Y')} | Planta CO01")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  ÍNDICE GENERAL (manual)
# ═══════════════════════════════════════════════════════════════════════════════
h1("Índice General")
secciones_indice = [
    ("1", "Resumen del Proyecto", "3"),
    ("2", "Arquitectura General", "4"),
    ("3", "Bases de Datos — Dependencias Completas", "5"),
    ("3.1", "BD SAP Azure (agpcolsap) — Solo lectura", "5"),
    ("3.2", "BD Producción Azure (agpcol) — Solo lectura", "7"),
    ("3.3", "BD Local SQL Express — Lectura/Escritura", "8"),
    ("4", "Rutas Flask (app.py)", "11"),
    ("5", "Funciones de Consulta (Helpers)", "14"),
    ("6", "Catálogos y Constantes", "17"),
    ("7", "Flujos SAP — Automatización", "18"),
    ("7.1", "Cambio de Color", "18"),
    ("7.2", "Cambio de Fórmula Sin Acero → Sin Acero", "20"),
    ("7.3", "Cambio de Fórmula Con Acero → Con Acero", "22"),
    ("7.4", "Cambio de Fórmula Mismo Acero (con CA02)", "23"),
    ("8", "Métodos SAP — Referencia Completa", "25"),
    ("9", "IDs SAP GUI Confirmados", "30"),
    ("10", "Pestaña Explorador", "33"),
    ("11", "Pestaña Hojas de Ruta", "34"),
    ("12", "Pestaña Cola de Homologaciones", "36"),
    ("13", "Pestaña Combinaciones", "38"),
    ("14", "Timings y Rendimiento", "39"),
    ("15", "Usuarios y Seguridad", "40"),
    ("16", "Estructura ResultadoItem", "41"),
]
t_indice = doc.add_table(rows=len(secciones_indice), cols=3)
t_indice.style = 'Table Grid'
for i, (num, titulo, pag) in enumerate(secciones_indice):
    t_indice.rows[i].cells[0].text = num
    t_indice.rows[i].cells[1].text = titulo
    t_indice.rows[i].cells[2].text = pag
    for ci in range(3):
        t_indice.rows[i].cells[ci].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  1. RESUMEN DEL PROYECTO
# ═══════════════════════════════════════════════════════════════════════════════
h1("1. Resumen del Proyecto")
para(
    "Este sistema es el corazón de la ingeniería de producto en AGP Glass Colombia. "
    "Su trabajo principal es automatizar todo lo que normalmente un ingeniero haría a mano en SAP: "
    "crear nuevas variantes de vidrios blindados (ZFERs), asignarles fórmula, color, acero, plano, "
    "hoja de ruta, subproducto y diferenciales. Todo eso que antes tomaba horas, ahora lo hace el sistema "
    "en minutos sin que nadie tenga que tocar SAP directo."
)
separador()
tabla(
    ["Campo", "Valor"],
    [
        ["Empresa", "AGP Glass — Planta Colombia CO01"],
        ["Entorno SAP Pruebas (QUAS)", "Usuario: PROGRAING / Pwd: AGPcol123*"],
        ["Tecnología principal", "Flask (Python) + SAP GUI Scripting (win32com)"],
        ["BD SAP (Azure)", "agpcolsap.database.windows.net / DB_COL_SAP"],
        ["BD Producción (Azure)", "agpcol.database.windows.net / agpc-productivity"],
        ["BD Local", r"localhost\SQLEXPRESS / MODULO_5"],
        ["Puerto Flask", "5000 (http://localhost:5000)"],
        ["Archivos principales", "app.py, sap_auto.py, templates/*.html"],
    ],
    col_widths=[5, 12]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  2. ARQUITECTURA GENERAL
# ═══════════════════════════════════════════════════════════════════════════════
h1("2. Arquitectura General")
para(
    "El sistema corre como una app web Flask en la máquina local del ingeniero. "
    "Tiene varias 'pestañas' (páginas web) que corresponden a distintas funcionalidades. "
    "Cuando se necesita ejecutar algo en SAP, Flask llama al automatizador (sap_auto.py) "
    "que controla el SAP GUI abierto en el mismo PC vía COM scripting."
)
separador()
h2("Pestañas (páginas web)")
tabla(
    ["URL", "Nombre", "Qué hace"],
    [
        ["/", "Login", "Autenticación de usuarios AGP"],
        ["/explorar", "Explorador de ZFERs", "Busca y muestra atributos de cualquier ZFER"],
        ["/zfer/<material>", "Detalle ZFER", "Ficha completa: atributos, BOM, entregas, variantes, simetría"],
        ["/combinaciones", "Combinaciones", "Genera combinaciones fórmula×color para un ZFER base"],
        ["/hojas_ruta", "Hojas de Ruta", "Busca la HR candidata y la asigna en SAP CA02"],
        ["/cola", "Cola SAP", "Gestiona la cola de homologaciones pendientes/programadas"],
        ["/api/*", "APIs internas", "Endpoints JSON para planos, atributos, ejecución SAP, etc."],
    ],
    col_widths=[4, 4, 9]
)
separador()
h2("Flujo de datos principal")
bullet("Usuario abre el explorador → Flask consulta BD SAP Azure → muestra datos")
bullet("Usuario envía un ZFER a la cola → Flask guarda en BD Local → worker SAP lo procesa")
bullet("Worker SAP usa sap_auto.py → controla SAP GUI → escribe resultados en BD Local")
bullet("Usuario ve el estado en /cola en tiempo real (polling cada 3s vía JS)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  3. BASES DE DATOS
# ═══════════════════════════════════════════════════════════════════════════════
h1("3. Bases de Datos — Dependencias Completas")
para(
    "Acá está TODO lo que toca el sistema a nivel de base de datos. "
    "Ni una tabla ni una columna se escapa. Organizado por servidor."
)

# ── 3.1 BD SAP Azure ──────────────────────────────────────────────────────────
h2("3.1 BD SAP Azure — agpcolsap.database.windows.net / DB_COL_SAP")
para("Solo lectura. Estas tablas replican datos de SAP en tiempo casi real. "
     "El sistema NUNCA escribe aquí.", italic=True)
separador()

h3("ODATA_ZFER_HEAD — Encabezado de materiales ZFER")
para("Es la tabla maestra de ZFERs. Contiene la existencia y estado de cada pieza.")
tabla(
    ["Columna", "Tipo", "Descripción", "Usada en"],
    [
        ["MATERIAL", "NVARCHAR", "Número SAP del ZFER (ej: 700179044)", "q_zfer_head, q_explorar, q_variantes_por_pn"],
        ["CENTRO", "NVARCHAR", "Centro productivo (siempre CO01)", "Todos los filtros WHERE"],
        ["TEXTO_BREVE_MATERIAL", "NVARCHAR", "Descripción corta del material", "Explorador, ficha ZFER"],
        ["STATUS", "NVARCHAR", "Estado SAP — ZZ = bloqueado/inactivo", "Filtro: UPPER(STATUS) != 'ZZ'"],
        ["ZFOR", "NVARCHAR", "ZFOR asociado (pieza interior)", "Ficha ZFER, cola SAP"],
        ["GRUPO_ARTICULOS", "NVARCHAR", "Grupo de artículos SAP", "Ficha ZFER"],
        ["CREADO_EL", "DATE", "Fecha de creación en SAP", "Ficha ZFER"],
        ["ULTIMA_MOD", "DATE", "Última modificación SAP", "Ficha ZFER"],
        ["AREA", "FLOAT", "Área de la pieza en m²", "Criterios HR: PEQUEÑA/MEDIANA/GRANDE"],
    ],
    col_widths=[4, 3, 6, 5]
)
separador()

h3("ODATA_ZFER_CLASS_001 — Atributos de clasificación ZFER")
para("Tabla pivote: cada fila es un atributo de un ZFER. "
     "Para leer todos los atributos de un ZFER se hace un GROUP BY con MAX(CASE WHEN...).")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["MATERIAL", "Número ZFER", "Join con ODATA_ZFER_HEAD"],
        ["CENTRO", "Centro CO01", "Filtro"],
        ["ATNAM", "Nombre del atributo SAP (ej: Z_COLOR)", "WHERE ATNAM IN (...)"],
        ["ATWRT", "Valor del atributo como texto", "La mayoría de consultas"],
        ["ATFLV", "Valor numérico (solo Z_COMMERCIAL_THICKNESS)", "CAST(ATFLV AS VARCHAR) para espesor"],
        ["TIPO_MAT", "Tipo material SAP", "Filtros ZPLA: TIPO_MAT='ZPLA'"],
    ],
    col_widths=[4, 8, 6]
)
separador()
para("Atributos (ATNAM) que lee el sistema:")
tabla(
    ["ATNAM", "Descripción", "Usada en función"],
    [
        ["Z_VEHICLE_MODEL", "Modelo del vehículo (ej: Toyota Hilux)", "q_atributos, q_explorar"],
        ["Z_AGP_PARTNUMBER", "Partnumber AGP (ej: 1490_008_L23-26_12_000)", "q_atributos, q_variantes_por_pn"],
        ["Z_FORMULA_CODE", "Código de fórmula (ej: L23-26)", "q_atributos, q_zplas_compatibles, cola"],
        ["Z_COLOR", "Código de color (ej: 19)", "q_atributos, q_zplas_compatibles"],
        ["Z_PIECE_TYPE", "Tipo de pieza (ej: 000 = Parabrisas)", "q_atributos, q_formulas_por_pieza, criterios HR"],
        ["Z_SHADE_BAND", "Franja (00/01/02/03/NA)", "q_atributos, franja en ZMME0001"],
        ["Z_AGP_LEVEL", "Nivel de protección AGP (1-5)", "q_atributos, criterios HR"],
        ["Z_BEHAVIOR_DIFFERENTIALS", "Diferenciales CSV (ej: 01,06,08)", "q_atributos, mm02_actualizar_diferenciales_zpla"],
        ["Z_SUBPRODUCT", "Código subproducto (ej: B3, X15)", "q_atributos, mm02_actualizar_subproducto"],
        ["Z_COMMERCIAL_THICKNESS", "Espesor comercial (numérico)", "Ficha ZFER display"],
        ["Z_AGP_VERSION", "Versión AGP del material", "q_explorar filtro, q_variantes_por_pn"],
        ["Z_GEOMETRY_TYPE", "Tipo geometría: 01=Plano, 02=Curvo", "Criterios HR: PLANO/CURVO"],
    ],
    col_widths=[5, 7, 6]
)
separador()

h3("ODATA_ZPLA_CLASS_001 — Atributos de clasificación ZPLA")
para("Igual que ODATA_ZFER_CLASS_001 pero para ZPLAs (materiales de referencia/plantilla). "
     "También tiene TIPO_MAT='ZPLA'.")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["MATERIAL", "Número ZPLA", "q_zplas_compatibles, q_formulas_por_pieza"],
        ["CENTRO", "Centro CO01", "Filtro"],
        ["ATNAM", "Nombre atributo", "WHERE ATNAM = 'Z_BEHAVIOR_DIFFERENTIALS'"],
        ["ATWRT", "Valor CSV (ej: '01,06,08')", "mm02_actualizar_diferenciales_zpla — _buscar_diferenciales_zpla"],
        ["TIPO_MAT", "ZPLA para filtrar solo plantillas", "_obtener_orden_diferenciales, _buscar_diferenciales_zpla"],
    ],
    col_widths=[4, 6, 8]
)
para("Consulta clave para diferenciales (sap_auto.py):", bold=True)
code(
    "SELECT DISTINCT value FROM ODATA_ZPLA_CLASS_001\n"
    "CROSS APPLY STRING_SPLIT(ATWRT, ',')\n"
    "WHERE ATNAM='Z_BEHAVIOR_DIFFERENTIALS' AND CENTRO='CO01' AND TIPO_MAT='ZPLA'\n"
    "ORDER BY value\n"
    "→ Retorna el orden de los diferenciales tal como aparecen en el popup SAP"
)
separador()

h3("ODATA_ZPLA_HEAD — Encabezado de materiales ZPLA")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["MATERIAL", "Número ZPLA", "q_zplas_compatibles, q_formulas_por_pieza — JOIN para filtrar STATUS != ZZ"],
        ["CENTRO", "Centro CO01", "Filtro"],
        ["STATUS", "Estado — NULL = activo, ZZ = inactivo", "WHERE STATUS IS NULL"],
    ],
    col_widths=[4, 6, 8]
)
separador()

h3("ODATA_ZFER_RUTAS_JPG — Planos técnicos")
para("Contiene las rutas UNC de los archivos JPG de los planos de cada ZFER.")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["MATERIAL", "Número ZFER", "WHERE MATERIAL=? AND CENTRO='CO01'"],
        ["CENTRO", "Centro CO01", "Filtro siempre presente"],
        ["PLANO", "Ruta UNC del archivo JPG (ej: \\\\servidor\\planos\\...)", "api_plano, mm02_cambiar_plano"],
        ["DOCUMENTO", "Nombre del documento SAP (ej: 'M1234 001 003 B')", "Filtro: NOT LIKE '% SP', NOT LIKE '% L[0-9]%'"],
        ["VERSION", "Versión del documento", "ORDER BY VERSION DESC para tomar la más reciente"],
        ["PROCESSDATE", "Fecha de proceso", "ORDER BY PROCESSDATE DESC (segundo criterio)"],
    ],
    col_widths=[3, 5, 4, 6]
)
para("Filtros SQL para búsqueda de plano (sap_auto.py):", bold=True)
code(
    "Sin SP:  AND DOCUMENTO NOT LIKE '% SP'\n"
    "         AND DOCUMENTO NOT LIKE '% L[0-9]%'    ← excluye documentos con código fórmula\n\n"
    "Con SP:  AND DOCUMENTO LIKE '% SP'\n"
    "         AND DOCUMENTO NOT LIKE '% L[0-9]%[0-9] SP'"
)
separador()

h3("ODATA_ZCDS_Entregas_Pos_CO — Posiciones de entregas")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["matnr", "Número de material ZFER", "WHERE matnr=?"],
        ["entrega", "Número de entrega SAP", "q_entregas → lista para q_mercados"],
        ["ntgew", "Peso neto", "WHERE TRY_CAST(ntgew AS FLOAT) > 0 (filtra líneas vacías)"],
    ],
    col_widths=[4, 6, 8]
)
separador()

h3("ODATA_ZCDS_Entregas_Head_CO — Cabecera de entregas")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["entrega", "Número de entrega", "WHERE entrega IN (lista de q_entregas)"],
        ["route", "Ruta/destino (código país ej: CO, AR, US)", "q_mercados → _decode_route → nombre país"],
    ],
    col_widths=[4, 6, 8]
)
separador()

h3("ODATA_HR_CONSULTA — Hojas de Ruta disponibles")
para("Tabla con todas las hojas de ruta de producción y sus criterios de asignación.")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["ID_HRUTA", "ID único de la hoja de ruta", "ca02_asignar_hr, c223_actualizar_version_fabricacion"],
        ["DESCRIPCION", "Nombre descriptivo de la HR", "Mostrado en /hojas_ruta UI"],
        ["TOTAL_MATERIALES", "Cuántos ZFERs ya usan esa HR", "Candidata: MAX donde <= 450"],
        ["NIVEL", "Nivel de protección (BAJO/MEDIO/ALTO)", "Criterio de filtrado en _hr_construir_criterios"],
        ["GEOMETRIA", "PLANO o CURVO", "Criterio geométrico"],
        ["FORMULA", "Código de fórmula compatible", "Criterio de fórmula"],
        ["TAMAÑO", "PEQUEÑA / MEDIANA / GRANDE", "Criterio por área"],
        ["SERIGRAFIA", "1 si la pieza tiene serigrafía", "excluir_null=True si criterio tiene valor"],
        ["MECANIZADO", "1 si requiere mecanizado", "excluir_null=True"],
        ["VITRIFICADO", "1 si requiere vitrificado", "excluir_null=True"],
        ["EMPALME", "Suma de posiciones para pieza curva", "Criterio empalme"],
        ["CURVADO", "1 si es curvo", "Criterio CURVO"],
        ["ENT_HORNO_CUR", "1 si entra al horno curvo", "Criterio"],
        ["SAL_HORNO_CUR", "1 si sale del horno curvo", "Criterio"],
        ["CURV_ACERO", "1 si tiene curvado con acero", "Criterio: pos 106/116 en BOM"],
        ["METROLOGIA", "1 si requiere metrología", "Heredado del ZFER base"],
        ["PRUEBA_AGUA", "1 si requiere prueba de agua", "Heredado del ZFER base"],
        ["BASE", "Posición base (32VPMO)", "Criterio posición 99 del BOM"],
        ["PROTECTORS", "Posiciones protectores", "Criterio pos 199/299"],
        ["TAPAS", "Posiciones tapas", "Criterio pos 3600/3700"],
    ],
    col_widths=[4, 6, 4, 4]
)

doc.add_page_break()

# ── 3.2 BD Producción Azure ───────────────────────────────────────────────────
h2("3.2 BD Producción Azure — agpcol.database.windows.net / agpc-productivity")
para("Solo lectura. Contiene datos comerciales y de clasificación adicionales.", italic=True)
separador()

h3("ODATA_ZPLA_CLASS_001 (también en esta BD)")
para("Esta BD también tiene esta tabla con los mismos atributos. "
     "Se usa específicamente en sap_auto.py para obtener los diferenciales del ZPLA base.")
code(
    "Función: _buscar_diferenciales_zpla(zpla_base)\n"
    "→ SELECT ATWRT FROM ODATA_ZPLA_CLASS_001\n"
    "  WHERE MATERIAL=zpla_base AND ATNAM='Z_BEHAVIOR_DIFFERENTIALS' AND CENTRO='CO01'"
)
separador()

h3("ZFER_Characteristics_Genesis — Fórmula del ZFER base")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["SpecID", "Número ZFER (equivale a MATERIAL)", "WHERE SpecID = zfer_base"],
        ["FormulaCode", "Código de fórmula (ej: L19-13)", "Pre-vuelo cola: detectar fórmula base"],
    ],
    col_widths=[4, 7, 7]
)
para("Es la fuente primaria para detectar la fórmula base. Si no encuentra aquí, va al fallback.", italic=True)
separador()

h3("TCAL_CALENDARIO_COLOMBIA_DIRECT — Fallback de fórmula")
tabla(
    ["Columna", "Descripción", "Usada en"],
    [
        ["ZFER", "Número ZFER", "WHERE ZFER = zfer_base (fallback si no está en Genesis)"],
        ["Formula", "Código de fórmula", "Pre-vuelo cola: segunda opción"],
        ["Mercado", "Mercado del ZFER", "Info adicional"],
        ["Color", "Color del ZFER", "Info adicional"],
    ],
    col_widths=[4, 6, 8]
)

doc.add_page_break()

# ── 3.3 BD Local ──────────────────────────────────────────────────────────────
h2("3.3 BD Local SQL Express — localhost\\SQLEXPRESS / MODULO_5")
para("Lectura y escritura. Aquí vive toda la lógica de la cola, los bloques y el historial.", italic=True)
separador()

h3("dbo.M5_Bloques — Bloques de homologación")
para("Un bloque es un grupo de ZFERs que se procesan juntos. "
     "Puede tener una hora programada para ejecutarse automáticamente.")
tabla(
    ["Columna", "Tipo", "Descripción"],
    [
        ["id", "INT PK", "ID único del bloque"],
        ["nombre", "NVARCHAR", "Nombre descriptivo del bloque"],
        ["estado", "NVARCHAR", "PENDIENTE / EJECUTANDO / COMPLETADO / ERROR"],
        ["timer_activo", "BIT", "1 si tiene hora programada activa"],
        ["hora_prog", "DATETIME", "Hora a la que debe ejecutarse automáticamente"],
        ["fecha_creacion", "DATETIME", "Cuándo se creó el bloque"],
        ["fecha_inicio", "DATETIME", "Cuándo empezó a ejecutarse"],
        ["fecha_fin", "DATETIME", "Cuándo terminó"],
        ["usuario", "NVARCHAR", "Email del usuario que lo creó"],
        ["total_items", "INT", "Total de items en el bloque"],
        ["items_ok", "INT", "Cuántos terminaron OK"],
        ["items_error", "INT", "Cuántos dieron error"],
    ],
    col_widths=[4, 3, 11]
)
separador()

h3("dbo.M5_Cola — Items de la cola (homologaciones individuales)")
para("Cada fila es una tarea SAP individual dentro de un bloque.")
tabla(
    ["Columna", "Tipo", "Descripción"],
    [
        ["id", "INT PK IDENTITY", "ID único del item"],
        ["bloque_id", "INT FK", "Referencia a M5_Bloques.id"],
        ["zfer_base", "NVARCHAR(20)", "ZFER base de entrada"],
        ["color_codigo", "NVARCHAR(5)", "Código del color SAP (ej: '19')"],
        ["color_nombre", "NVARCHAR(100)", "Nombre descriptivo del color"],
        ["zpla", "NVARCHAR(20)", "ZPLA de referencia para el proceso"],
        ["franja", "NVARCHAR(5)", "Código de franja (00/01/02/03/NA)"],
        ["tipo", "NVARCHAR(20)", "color / FORMULA_SIN_ACERO / FORMULA_CON_ACERO / FORMULA_MISMO_ACERO"],
        ["formula_nueva", "NVARCHAR(20)", "Nueva fórmula (solo cambios de fórmula)"],
        ["acero_dir", "NVARCHAR(20)", "Dirección del acero: con_sin / sin_con / mismo"],
        ["cambiar_hr", "BIT", "1 si se debe asignar nueva hoja de ruta"],
        ["estado", "NVARCHAR(20)", "PENDIENTE / EJECUTANDO / COMPLETADO / ERROR"],
        ["zfer_nuevo", "NVARCHAR(20)", "ZFER creado por SAP (se llena al completar)"],
        ["zfor_nuevo", "NVARCHAR(20)", "ZFOR creado (si aplica)"],
        ["error_msg", "NVARCHAR(MAX)", "Mensaje de error si falló"],
        ["advertencias", "NVARCHAR(MAX)", "Advertencias no fatales (plano no encontrado, etc.)"],
        ["fecha_inicio", "DATETIME", "Inicio del procesamiento SAP"],
        ["fecha_fin", "DATETIME", "Fin del procesamiento SAP"],
        ["subproducto", "NVARCHAR(20)", "Subproducto del ZFER nuevo (NULL = no aplica)"],
        ["tipo_pieza", "NVARCHAR(10)", "Tipo de pieza para referencia"],
        ["formula", "NVARCHAR(20)", "Fórmula del ZFER base"],
        ["acero", "NVARCHAR(100)", "Descripción del tipo de acero"],
    ],
    col_widths=[4, 4, 10]
)
separador()

h3("dbo.M5_LogEjecucion — Log histórico de ejecuciones")
para("Historial permanente. Cada ZFER procesado deja un registro aquí. "
     "Si la BD local no existe, los errores se ignoran (solo warning en consola).")
tabla(
    ["Columna", "Tipo", "Descripción"],
    [
        ["id", "INT PK IDENTITY", "ID autoincremental"],
        ["batch_id", "VARCHAR(50)", "UUID del lote de procesamiento"],
        ["zfer_base", "NVARCHAR(20)", "ZFER de entrada"],
        ["zfer_nuevo", "NVARCHAR(20)", "ZFER creado por SAP"],
        ["zfor_nuevo", "NVARCHAR(20)", "ZFOR creado (si aplica)"],
        ["zpla", "NVARCHAR(20)", "ZPLA usado"],
        ["color_codigo", "NVARCHAR(5)", "Código color"],
        ["color_nombre", "NVARCHAR(100)", "Nombre color"],
        ["formula", "NVARCHAR(20)", "Fórmula del ZFER base"],
        ["tipo_pieza", "NVARCHAR(10)", "Tipo de pieza"],
        ["acero", "NVARCHAR(100)", "Tipo de acero"],
        ["estado", "NVARCHAR(20)", "OK / ERROR"],
        ["error_msg", "NVARCHAR(MAX)", "Detalle del error si aplica"],
        ["duracion_seg", "FLOAT", "Tiempo de procesamiento en segundos"],
        ["fecha", "DATETIME", "Fecha y hora de ejecución"],
        ["tipo", "NVARCHAR(20)", "color / formula"],
    ],
    col_widths=[4, 4, 10]
)
separador()

h3("dbo.M5_LogEjecuciones — Log extendido (variante con más campos)")
para("Versión extendida del log que incluye advertencias y posiciones BOM.")
tabla(
    ["Columna extra vs M5_LogEjecucion", "Descripción"],
    [
        ["advertencias", "Texto de advertencias no fatales (plano, HR, etc.)"],
        ["posiciones_bom", "Posiciones BOM procesadas en formato JSON/CSV"],
    ],
    col_widths=[6, 12]
)
separador()

h3("dbo.M5_HomologacionFormula — Registro de cambios de fórmula")
tabla(
    ["Columna", "Tipo", "Descripción"],
    [
        ["id", "INT PK", "ID único"],
        ["zfer_base", "NVARCHAR(20)", "ZFER de entrada"],
        ["formula_base", "NVARCHAR(20)", "Fórmula del ZFER base"],
        ["formula_nueva", "NVARCHAR(20)", "Nueva fórmula objetivo"],
        ["tipo", "NVARCHAR(30)", "FORMULA_SIN_ACERO / FORMULA_CON_ACERO / FORMULA_MISMO_ACERO"],
        ["estado", "NVARCHAR(20)", "PENDIENTE / OK / ERROR"],
        ["fecha", "DATETIME", "Fecha de creación del registro"],
        ["usuario", "NVARCHAR(100)", "Usuario que lo solicitó"],
    ],
    col_widths=[4, 4, 10]
)
separador()

h3("dbo.M5_HomologacionFormula_BOM — BOM de cambios de fórmula")
para("Detalle de las posiciones BOM procesadas para cada cambio de fórmula.")
tabla(
    ["Columna", "Descripción"],
    [
        ["id_homologacion", "FK a M5_HomologacionFormula.id"],
        ["posnr", "Número de posición BOM (ej: 0458)"],
        ["clase_destino", "Clase destino consultada en ODATA_ZPLA_BOM"],
        ["estado", "OK / ERROR por posición"],
    ],
    col_widths=[5, 13]
)
separador()

h3("dbo.M5_RutasZFER — Rutas asignadas a ZFERs")
para("Registro de asignaciones de hojas de ruta.")
tabla(
    ["Columna", "Descripción"],
    [
        ["zfer", "Número ZFER"],
        ["id_hruta", "ID de la hoja de ruta asignada"],
        ["fecha", "Fecha de asignación"],
        ["usuario", "Usuario que hizo la asignación"],
        ["estado", "OK / ERROR"],
    ],
    col_widths=[4, 6, 8]
)
separador()

h3("dbo.M5_Bloqueos — Combinaciones bloqueadas")
para("Combinaciones que el usuario marcó como bloqueadas en la UI.")
tabla(
    ["Columna", "Descripción"],
    [
        ["zfer_base", "ZFER base"],
        ["formula", "Código de fórmula"],
        ["color", "Código de color"],
        ["acero", "Tipo de acero"],
        ["motivo", "Razón del bloqueo (texto libre)"],
        ["fecha", "Fecha del bloqueo"],
        ["usuario", "Usuario que bloqueó"],
    ],
    col_widths=[4, 5, 9]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  4. RUTAS FLASK
# ═══════════════════════════════════════════════════════════════════════════════
h1("4. Rutas Flask (app.py) — Referencia Completa")
para("Todas las rutas web del sistema. Las que empiezan con /api/ retornan JSON.")
separador()

h2("Rutas de autenticación")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["GET/POST", "/login", "login()", "Formulario de login. Valida contra dict _USUARIOS"],
        ["GET", "/logout", "logout()", "Cierra la sesión Flask"],
    ],
    col_widths=[2, 4, 4, 8]
)
separador()

h2("Rutas del Explorador")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["GET", "/explorar", "explorar()", "Página principal del explorador. Acepta query params: vehiculo, formula, pieza, color, version, nivel, cod_vehiculo"],
        ["GET", "/zfer/<material>", "zfer_detalle(material)", "Ficha completa de un ZFER: atributos, BOM SAP, variantes de color, entregas, mercados, plano"],
        ["GET", "/api/zfer/<material>/attrs", "api_zfer_attrs()", "JSON con atributos de clasificación del ZFER (q_atributos)"],
        ["GET", "/api/plano/<material>", "api_plano(material)", "Sirve el archivo JPG del plano directamente (Content-Type: image/jpeg)"],
        ["GET", "/api/planos/batch", "api_planos_batch()", "JSON {material: documento} para múltiples materiales en un solo query (param: mats=M1,M2,M3)"],
    ],
    col_widths=[2, 5, 5, 6]
)
separador()

h2("Rutas de Combinaciones")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["GET/POST", "/combinaciones", "combinaciones()", "Página de generación de combinaciones fórmula×color para un ZFER base"],
        ["GET", "/api/combinaciones/<zfer>", "api_combinaciones(zfer)", "JSON con todas las combinaciones posibles para el ZFER"],
        ["POST", "/api/combinaciones/enviar_cola", "api_enviar_cola()", "Envía combinaciones seleccionadas a la cola SAP"],
    ],
    col_widths=[2, 5, 5, 6]
)
separador()

h2("Rutas de Cola SAP")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["GET", "/cola", "cola()", "Página principal de la cola. Muestra bloques y sus items"],
        ["POST", "/api/cola/crear_bloque", "api_crear_bloque()", "Crea un nuevo bloque en M5_Bloques"],
        ["POST", "/api/cola/agregar_items", "api_agregar_items()", "Agrega items a un bloque existente en M5_Cola"],
        ["POST", "/api/cola/ejecutar/<bloque_id>", "api_ejecutar_bloque()", "Dispara la ejecución inmediata de un bloque"],
        ["POST", "/api/cola/programar/<bloque_id>", "api_programar_bloque()", "Programa un bloque para ejecutarse a una hora específica"],
        ["GET", "/api/cola/estado", "api_estado_cola()", "JSON con el estado actual de todos los bloques y sus items (polling cada 3s desde UI)"],
        ["GET", "/api/cola/detalle/<bloque_id>", "api_detalle_bloque()", "JSON con el detalle completo de un bloque (items, logs, advertencias)"],
        ["POST", "/api/cola/cancelar/<bloque_id>", "api_cancelar_bloque()", "Cancela un bloque PENDIENTE"],
        ["DELETE", "/api/cola/eliminar/<bloque_id>", "api_eliminar_bloque()", "Elimina un bloque y todos sus items de la BD"],
        ["GET", "/api/cola/descargar/<bloque_id>", "api_descargar_reporte()", "Genera y descarga el reporte Excel del bloque (4 hojas)"],
    ],
    col_widths=[2, 5, 5, 6]
)
separador()

h2("Rutas de Hojas de Ruta")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["GET", "/hojas_ruta", "hojas_ruta()", "Página de búsqueda y asignación de hojas de ruta. Params: zfer_base, zfer_nuevo"],
        ["GET", "/api/hr/candidata", "api_hr_candidata()", "JSON con la HR candidata para un ZFER: id_hruta, descripcion, materiales, criterios usados"],
        ["GET", "/api/hr/todas", "api_hr_todas()", "JSON con todas las HRs que cumplen criterios (para tab 'Otras opciones')"],
        ["POST", "/api/hr/asignar", "api_hr_asignar()", "Ejecuta CA02 + C223 en SAP para asignar la HR al ZFER"],
        ["GET", "/api/hr/bom/<zfer>", "api_hr_bom(zfer)", "JSON con el BOM SAP del ZFER (posiciones y materiales)"],
        ["GET", "/api/hr/criterios", "api_hr_criterios()", "JSON con los criterios construidos para diagnóstico (tab SQL/Criterios)"],
    ],
    col_widths=[2, 5, 5, 6]
)
separador()

h2("Rutas SAP directas")
tabla(
    ["Método", "Ruta", "Función", "Descripción"],
    [
        ["POST", "/api/sap/procesar", "api_sap_procesar()", "Inicia procesamiento SAP de un item (lanza hilo). Retorna batch_id"],
        ["GET", "/api/sap/estado/<batch_id>", "api_sap_estado(batch_id)", "Polling del estado de un job SAP activo"],
        ["GET", "/api/sap/log/<batch_id>", "api_sap_log(batch_id)", "Retorna el log completo de un job SAP"],
    ],
    col_widths=[2, 5, 5, 6]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  5. FUNCIONES DE CONSULTA
# ═══════════════════════════════════════════════════════════════════════════════
h1("5. Funciones de Consulta (Helpers) — app.py")
para(
    "Estas son las funciones que hacen las consultas a la BD y procesan los datos "
    "para mostrárselos al usuario o pasárselos al automatizador SAP."
)
separador()

h2("Conexión y Pool")
tabla(
    ["Función", "Descripción", "Tablas BD"],
    [
        ["get_conn()", "Obtiene conexión del pool (12 conexiones máx). Si el pool está vacío crea una nueva. Health-check con SELECT 1 antes de usar.", "DB_COL_SAP (agpcolsap)"],
        ["_get_conn_local()", "Conexión a la BD local SQL Express. Sin pool (Trusted_Connection).", "MODULO_5 (localhost\\SQLEXPRESS)"],
        ["_conn_str()", "Genera el string de conexión para DB_COL_SAP.", "—"],
    ],
    col_widths=[4, 10, 4]
)
separador()

h2("Consultas principales")
tabla(
    ["Función", "Retorna", "Tablas que toca", "Cacheada"],
    [
        ["q_zfer_head(material)", "dict con info básica del ZFER o None", "ODATA_ZFER_HEAD", "lru_cache(400)"],
        ["q_atributos(material)", "dict {ATNAM: valor} con todos los atributos de clasificación", "ODATA_ZFER_CLASS_001", "lru_cache(400)"],
        ["q_entregas(material)", "lista de números de entrega donde ntgew>0", "ODATA_ZCDS_Entregas_Pos_CO", "No"],
        ["q_mercados(entregas)", "lista [{route, pais, total}] — mercados por volumen", "ODATA_ZCDS_Entregas_Head_CO", "No"],
        ["q_variantes_por_pn(vehiculo, version, formula, pieza)", "lista de ZFERs con el mismo vehículo/fórmula/pieza en todos los colores", "ODATA_ZFER_CLASS_001 + ODATA_ZFER_HEAD (JOIN)", "No"],
        ["q_zplas_compatibles(formula_code, piece_type, shade_band, differentials_base, tiene_acero_base)", "lista de ZPLAs que coinciden con los criterios del ZFER base", "ODATA_ZPLA_CLASS_001 + ODATA_ZPLA_HEAD (JOIN)", "No"],
        ["q_formulas_por_pieza(piece_type, nivel, subproducto, formula_base)", "lista [{formula, colores:[{zpla, color, differentials}]}] — fórmulas alternativas disponibles", "ODATA_ZPLA_HEAD + ODATA_ZPLA_CLASS_001 (CTE)", "No"],
        ["q_explorar(...filtros...)", "lista de hasta 300 ZFERs que coinciden con los filtros de búsqueda", "ODATA_ZFER_CLASS_001 + ODATA_ZFER_HEAD (JOIN)", "No"],
    ],
    col_widths=[5, 5, 5, 3]
)
separador()

h2("Helpers de planos")
tabla(
    ["Función", "Descripción", "Tabla BD"],
    [
        ["_q_plano(material)", "Retorna (ruta_unc, documento) o None. Cacheado en _plano_cache dict.", "ODATA_ZFER_RUTAS_JPG"],
        ["_q_planos_bulk(mats)", "Un solo IN query para poblar _plano_cache con lista de materiales. Evita N queries individuales.", "ODATA_ZFER_RUTAS_JPG"],
        ["_normalizar_unc(ruta)", "Garantiza que la ruta empiece con \\\\\\\\ (UNC válido Windows).", "—"],
        ["_plano_base(documento)", "Extrae el nombre base del documento quitando versión letra y SP al final.", "—"],
    ],
    col_widths=[5, 10, 4]
)
separador()

h2("Helpers de Hojas de Ruta")
tabla(
    ["Función", "Descripción", "Tablas BD"],
    [
        ["_hr_construir_criterios(zfer_base, zfer_nuevo)", "Construye el dict de criterios para filtrar HRs: NIVEL, GEOMETRIA, TAMAÑO, FORMULA, EMPALME, CURVADO, etc.", "ODATA_ZFER_HEAD, ODATA_ZFER_CLASS_001, ODATA_HR_CONSULTA (vía leer_bom_material SAP)"],
        ["_hr_buscar_candidata(zfer_base, zfer_nuevo)", "Retorna (id_hruta, descripcion, error). La candidata es la HR con MAX(TOTAL_MATERIALES) donde TOTAL_MATERIALES <= 450.", "ODATA_HR_CONSULTA"],
        ["_hr_buscar_todas(zfer_base, zfer_nuevo)", "Igual pero retorna todas las HRs que cumplen criterios (para tab 'Otras opciones').", "ODATA_HR_CONSULTA"],
    ],
    col_widths=[5, 9, 5]
)
separador()

h2("Helpers de Cola")
tabla(
    ["Función", "Descripción"],
    [
        ["_cola_ejecutar_bloque(bloque_id)", "Worker principal: itera todos los items PENDIENTE del bloque y ejecuta cada uno con el tipo correcto de flujo SAP"],
        ["_cola_ejecutar_item(item, sap)", "Ejecuta un item individual: llama procesar_combinacion / procesar_formula_sin_acero / etc. según item['tipo']"],
        ["_cola_scheduler()", "Hilo daemon que cada 20s revisa M5_Bloques buscando bloques PENDIENTE con timer_activo=1 cuya hora_prog ya pasó"],
        ["_cola_scheduler_tick()", "Un ciclo del scheduler: consulta BD, detecta vencidos, dispara en hilos. Usa _scheduler_disparados set para dedup"],
        ["_cola_limpiar_al_inicio()", "Al iniciar Flask: resetea items EJECUTANDO → PENDIENTE (recovery de crash)"],
        ["_migracion_bd_local()", "Al iniciar Flask: agrega columnas nuevas a tablas existentes (ALTER TABLE ADD si no existe). Migración automática."],
    ],
    col_widths=[5, 13]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  6. CATÁLOGOS Y CONSTANTES
# ═══════════════════════════════════════════════════════════════════════════════
h1("6. Catálogos y Constantes")
para("Diccionarios Python definidos en app.py que mapean códigos SAP a nombres legibles.")
separador()

h2("PIEZAS — Tipos de piezas")
para("Dict de ~80 entradas: {'000': 'Parabrisas', '001': 'Lateral Delantero Izquierdo', ...}")
para("También incluye grupos generados: Piezas Especiales (041-059), Vidrios Especiales (060-069), Pieza Plana Especial (070-073), Vidrio Especial Laminado (080-086).")
separador()

h2("COLORES — Códigos de color")
tabla(
    ["Código", "Nombre", "¿Activo?"],
    [
        ["00", "Blanco", "✓ Activo"],
        ["01", "Green Light", "✓ Activo"],
        ["05", "Gray Light PC", "✓ Activo"],
        ["06", "Gray Light Glass", "✓ Activo"],
        ["10", "Gray Medium PC", "✓ Activo"],
        ["13", "Gray Dark Glass", "✓ Activo"],
        ["18", "Gray Medium Glass", "✓ Activo"],
        ["19", "Gray Light Automotive", "✓ Activo"],
        ["20", "Gray Medium Automotive + PC", "✓ Activo"],
        ["21", "Gray Dark Automotive + PC", "✓ Activo"],
        ["22", "G2 Gray Medium Automotive", "✓ Activo"],
        ["23", "G2 Gray Dark Automotive", "✓ Activo"],
        ["02-04, 07-09, 11-12, 14-17", "Otros colores", "No activos (no se muestran en combinaciones)"],
    ],
    col_widths=[3, 8, 3]
)
separador()

h2("DIFERENCIALES — Comportamientos de protección")
para("24 diferenciales (01-24). Los más importantes en el flujo SAP:")
tabla(
    ["Código", "Nombre", "Relevancia especial"],
    [
        ["01", "SOLAR PLUS", "Diferencial estándar"],
        ["06", "STEEL PLUS", "⚠️ Indica presencia de ACERO — crítico para detectar tipo de flujo"],
        ["07", "TNT", "Estándar"],
        ["08", "TNT FLEX", "Estándar"],
        ["23", "N.A", "Sin diferencial aplicable"],
    ],
    col_widths=[2, 6, 10]
)
separador()

h2("FRANJAS")
para("{'00': 'Sin Franja', '01': 'Franja Azul', '02': 'Franja Verde', '03': 'Franja Gris', 'NA': 'No Aplica'}")
para("Se lee del ZFER base y se pasa tal cual al campo P_FRANJ de ZMME0001.")
separador()

h2("_PARES_SIMETRIA")
para("Mapa de piezas simétricas izquierda↔derecha. Si se procesa la 001 (Lateral Der. Izq.), "
     "automáticamente se ofrece procesar la 002 (Lateral Der. Der.) como par simétrico.")
para("Pares: 001↔002, 003↔004, 005↔006, 007↔008, 011↔012, 013↔014, 015↔016, 019↔020, 021↔022, 023↔024, 026↔027, 028↔029")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  7. FLUJOS SAP
# ═══════════════════════════════════════════════════════════════════════════════
h1("7. Flujos SAP — Automatización")
para(
    "Acá está el corazón del asunto. Cada tipo de homologación tiene su propio flujo "
    "de pasos en SAP. Todos viven en sap_auto.py como métodos de la clase AutomatizadorSAP."
)
separador()

# ── 7.1 Cambio de Color ───────────────────────────────────────────────────────
h2("7.1 Flujo: Cambio de Color (procesar_combinacion)")
para("El flujo original. Toma un ZFER existente y crea una variante de color diferente.")
separador()
tabla(
    ["Paso", "Transacción SAP", "Método Python", "Descripción"],
    [
        ["1", "ZPPR0008", "zppr0008_validar_posicion_acero(zfer_base)", "Verifica si el ZFER tiene posición 0106 o 0116 en el BOM → detecta si tiene acero"],
        ["2", "ZMME0001", "zmme0001_ejecutar(zfer_base, color, franja, zplas_validos)", "Homologar → Cambio de Color → F4 ZPLA → F8 → obtiene ZFER_NUEVO y ZFOR_NUEVO"],
        ["3", "ZPPR0020", "zppr0020_esperar_fases(zfer_nuevo)", "Espera en sesión auxiliar hasta que las 8+ fases tengan estado 'S'. Máx 300 seg (60 intentos cada 5s)"],
        ["4", "ZMME0001", "zmme0001_leer_posiciones_popup() + zmme0001_agregar_filas_bom() + zmme0001_segunda_comparar_y_copy()", "Vuelve a ZMME0001 con ZFER_NUEVO → Comparar BOM → agrega posiciones → COPY_ITEM"],
        ["5a", "MM02", "mm02_actualizar_partnumber(zfer_nuevo, nuevo_pn)", "Actualiza el PARTNUMBER en la clasificación PIEZA"],
        ["5b", "MM02", "mm02_actualizar_subproducto(zfer_nuevo, subproducto)", "Actualiza el subproducto en la clasificación"],
        ["5c", "MM02", "mm02_actualizar_diferenciales_zpla(zfer_nuevo, zpla_base)", "Marca/desmarca los diferenciales según el ZPLA de referencia"],
        ["5d", "MM02", "mm02_cambiar_plano(zfer_nuevo, res)", "Asigna el plano técnico (sin SP). Verifica statusbar SAP al guardar"],
        ["6 (si cambiar_hr)", "CA02 + C223", "cambiar_hoja_ruta(zfer_nuevo, id_hruta)", "Desasigna HR anterior → asigna nueva → si CA02 OK: actualiza versión fabricación en C223"],
    ],
    col_widths=[1, 3, 6, 8]
)
separador()

# ── 7.2 Fórmula sin→sin ───────────────────────────────────────────────────────
h2("7.2 Flujo: Cambio de Fórmula Sin Acero → Sin Acero (procesar_formula_sin_acero)")
para("Para cuando la fórmula cambia pero la pieza no tiene acero en ninguno de los dos casos.")
separador()
tabla(
    ["Paso", "Transacción", "Método", "Descripción"],
    [
        ["1", "ZPPR0008", "zppr0008_validar_posicion_acero(zfer_base)", "Valida que NO tenga posición 0106/0116 (confirma que es sin acero)"],
        ["2", "ZMME0001", "zmme0001_cambio_formula(zfer_base, formula_nueva, zpla)", "Homologar → Cambio Fórmula (radRB2_A1) → P_FORMU → F8 → ZFER_NUEVO"],
        ["3", "ZPPR0020", "zppr0020_esperar_fases(zfer_nuevo)", "Polling fases, máx 300 seg"],
        ["4", "ZMME0001", "zmme0001_leer_posiciones_popup() + agregar + copy", "BOM con retry (hasta 3 ciclos)"],
        ["5a", "MM02", "mm02_actualizar_partnumber()", "PN con nuevo código de fórmula"],
        ["5b", "MM02", "mm02_actualizar_subproducto()", "Subproducto del ZPLA nuevo"],
        ["5c", "MM02", "mm02_actualizar_diferenciales_zpla()", "Diferenciales según ZPLA nuevo (SIN diferencial 06)"],
        ["5d", "MM02", "mm02_cambiar_plano()", "Plano SIN SP"],
        ["6", "CEWB", "cewb_eliminar_posicion_acero(zfer_nuevo)", "Elimina posición 0106/0116 si quedó del original"],
        ["7", "ZMME0001", "_volver_zmme0001()", "Vuelve a ZMME0001 para dejar la pantalla limpia"],
        ["8 (si cambiar_hr)", "CA02+C223", "cambiar_hoja_ruta()", "Asigna nueva HR"],
    ],
    col_widths=[1, 3, 5, 9]
)
separador()

# ── 7.3 Fórmula con→con ───────────────────────────────────────────────────────
h2("7.3 Flujo: Cambio de Fórmula Con Acero → Con Acero (procesar_formula_con_acero)")
para("Igual que sin→sin pero el plano es CON SP y en CEWB se agrega (no elimina) la posición 0116.")
separador()
tabla(
    ["Diferencia vs flujo sin acero", "Detalle"],
    [
        ["Paso 1", "ZPPR0008 valida que SÍ tenga posición 0106 o 0116"],
        ["Paso 5c", "mm02_actualizar_diferenciales_zpla() marca diferencial 06 (STEEL PLUS)"],
        ["Paso 5d", "mm02_cambiar_plano_con_sp() busca plano CON SP (LIKE '% SP')"],
        ["Paso 6", "cewb_agregar_posicion_acero(zfer_nuevo) AGREGA posición 0116 en lugar de eliminarla"],
    ],
    col_widths=[4, 14]
)
separador()

# ── 7.4 Fórmula mismo acero ───────────────────────────────────────────────────
h2("7.4 Flujo: Cambio de Fórmula Mismo Acero (procesar_formula_mismo_acero)")
para("Para cambios de fórmula donde el acero no cambia. Incluye opcionalmente CA02 y C223.")
separador()
tabla(
    ["Diferencia vs los otros flujos", "Detalle"],
    [
        ["Sin cambio en CEWB", "No se agrega ni elimina posición de acero"],
        ["CA02 es opcional", "Solo se ejecuta si cambiar_hr=True en el item de cola"],
        ["C223 depende de CA02", "c223_actualizar_version_fabricacion() SOLO si ca02_asignar_hr() retorna True"],
        ["Plano con o sin SP", "Depende de si el ZFER tiene acero: mm02_cambiar_plano() o mm02_cambiar_plano_con_sp()"],
    ],
    col_widths=[5, 13]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  8. MÉTODOS SAP — REFERENCIA COMPLETA
# ═══════════════════════════════════════════════════════════════════════════════
h1("8. Métodos SAP — Referencia Completa (sap_auto.py)")
para("Todos los métodos de la clase AutomatizadorSAP, ordenados por transacción SAP.")
separador()

h2("Conexión")
tabla(
    ["Método", "Descripción"],
    [
        ["conectar()", "Conecta vía COM a SAP GUI activo. Obtiene app → conn_sap → session. Maximiza ventana."],
    ],
    col_widths=[5, 13]
)
separador()

h2("Helpers internos")
tabla(
    ["Método", "Descripción"],
    [
        ["_esperar(max_seg)", "Espera inteligente: mínimo garantizado (0.02/0.03/0.05s) + poll de app.Busy cada 50ms hasta max_seg"],
        ["_navegar(tcode)", "Escribe /N{tcode} en el campo de T-Code y presiona Enter"],
        ["_sbar()", "Lee el statusbar SAP. Retorna (tipo, texto). tipo: 'E'=error, 'W'=warning, 'S'=success, 'I'=info"],
        ["_estado_sap()", "Retorna solo el texto del statusbar"],
        ["_aceptar_dialogo()", "Envía Enter a wnd[1] para cerrar diálogos simples"],
        ["_cerrar_dialogs_abiertos()", "Cierra wnd[2] y wnd[1] si están abiertos (con F12 o Enter)"],
    ],
    col_widths=[5, 13]
)
separador()

h2("ZPPP0042 — Validación de versión")
tabla(
    ["Método", "Descripción"],
    [
        ["zppp0042_validar(zfer)", "Navega a ZPPP0042, busca el ZFER en el grid y verifica que la utilización = '1'. Retorna {ok, error, verid}"],
    ],
    col_widths=[5, 13]
)
separador()

h2("ZPPR0008 — BOM / Posiciones de acero")
tabla(
    ["Método", "Descripción"],
    [
        ["zppr0008_validar_posicion_acero(zfer)", "Lee el BOM del ZFER y verifica si tiene posición 0106 o 0116. Retorna {tiene_acero, posicion}"],
        ["leer_bom_material(zfer)", "Lee el BOM completo del ZFER. Retorna lista de {posnr, material, cantidad, unidad}. Usado también por /hojas_ruta"],
    ],
    col_widths=[6, 12]
)
separador()

h2("ZMME0001 — Homologaciones")
tabla(
    ["Método", "Descripción"],
    [
        ["zmme0001_ejecutar(zfer_base, color, franja, zplas_validos, forzar_be)", "Flujo de Cambio de Color: Homologar → F4 ZPLA → valida contra zplas_validos → doble clic → F8. Retorna (zfer_nuevo, zfor_nuevo, zpla_sel)"],
        ["zmme0001_cambio_formula(zfer_base, formula_nueva, zpla)", "Flujo de Cambio de Fórmula: usa radRB2_A1 y campo P_FORMU. Retorna (zfer_nuevo, zfor_nuevo)"],
        ["zmme0001_leer_posiciones_popup()", "Presiona Comparar BOM (btnBUTTON1) → lee popup tblZMME0001T_COMP → retorna lista de posiciones (ej: ['0458'])"],
        ["zmme0001_agregar_filas_bom(posiciones, zpla)", "Por cada posición: Insert → llena POSNR y CLASE_DESTINO (consultada en ODATA_ZPLA_BOM según el ZPLA)"],
        ["zmme0001_segunda_comparar_y_copy()", "Segunda Comparar BOM → verifica no hay error en popup → presiona COPY_ITEM"],
    ],
    col_widths=[6, 12]
)
separador()

h2("ZPPR0020 — Polling de fases")
tabla(
    ["Método", "Descripción"],
    [
        ["zppr0020_esperar_fases(zfer_nuevo, max_espera_seg=300)", "Navega a ZPPR0020 en sesión auxiliar. Polling cada 5s (60 intentos): busca la fila del ZFER en el grid ALV y verifica que fase 8+ tenga estado 'S'. Si alguna fase tiene 'E' → aborta con error. Retorna {ok, zpla, fase_error, detalle, fases}"],
    ],
    col_widths=[6, 12]
)
separador()

h2("MM02 — Actualización de materiales")
tabla(
    ["Método", "Descripción"],
    [
        ["mm02_actualizar_partnumber(zfer, nuevo_pn)", "Navega a MM02 → tab Clasificación → tab PIEZA → actualiza fila 0 (PARTNUMBER). Formato: vehiculo_seq_formula_color_pieza"],
        ["mm02_actualizar_subproducto(zfer, subproducto)", "scroll=3, vis_row=8 → escribe subproducto → sendVKey(0) → guarda. Toca tblSAPLCTMSCHARS_S"],
        ["mm02_actualizar_diferenciales_zpla(zfer, zpla_base)", "Abre popup Z_BEHAVIOR_DIFFERENTIALS → marca/desmarca checkboxes según lo que tenga el ZPLA base. Usa paginación por bloques (vis_pop=10), sin setFocus, 150ms entre páginas"],
        ["mm02_cambiar_plano(zfer, res)", "Busca plano SIN SP en ODATA_ZFER_RUTAS_JPG → actualiza DMS en MM02. Verifica statusbar: error E → advertencia en reporte"],
        ["mm02_cambiar_plano_con_sp(zfer, res)", "Igual pero busca plano CON SP (LIKE '% SP')"],
        ["_mm02_navegar_pieza_tab(zfer)", "Navega a MM02 → abre el material → va al tab Clasificación → sub-tab PIEZA. Helper compartido."],
        ["_mm02_guardar_y_salir()", "Presiona guardar (btn[11]) → confirma popup si aparece → vuelve (btn[15])"],
        ["_mm02_buscar_fila_car(tbl, atnam)", "Busca en la tabla de clasificación la fila que corresponde al atributo ATNAM dado"],
    ],
    col_widths=[6, 12]
)
separador()

h2("CA02 — Hojas de Ruta")
tabla(
    ["Método", "Descripción"],
    [
        ["ca02_desasignar_hr(zfer_nuevo)", "Navega a CA02 → abre popup de materiales → escanea con block-skip buscando el ZFER → selecciona fila → btn[14] borrar → confirma. Si no tiene HR: warning y continúa"],
        ["ca02_asignar_hr(zfer_nuevo, id_hruta)", "Igual pero al revés: abre HR por ID → popup materiales → scroll al fondo → find fila vacía → modifyCell con ZFER y CO01 → btn[0] confirmar → btn[11] guardar. Retorna True si OK"],
        ["_ca02_scroll(tbl, pos)", "Mueve el scrollbar de la tabla CA02 a la posición dada. time.sleep(0.02) después (sin espera SAP)"],
        ["_ca02_leer_matnr_vis(tbl_id, vis_row)", "Lee el valor de la columna MATNR[2,vis_row] usando findById. Retorna '' si error"],
    ],
    col_widths=[5, 13]
)
separador()

h2("C223 — Versión de Fabricación")
tabla(
    ["Método", "Descripción"],
    [
        ["c223_actualizar_version_fabricacion(zfer_nuevo, id_hruta, res)", "Navega a C223. Limpia TODOS los campos de filtro (PLNTY, PLNNR, BEDPL, DATUV, FERTH). Escribe CO01 + zfer_nuevo → Enter → escribe id_hruta en ctxtMKAL_EXPAND-PLNNR[16,0] → Enter → confirma popup → foco MATNR[1,0] → VKey 2 → btnPRUEFEN → btn[12] → btn[8] → guarda btn[11] → sale btn[15]. Solo se llama si CA02 retorna True"],
    ],
    col_widths=[6, 12]
)
separador()

h2("CEWB — Estructura de producto")
tabla(
    ["Método", "Descripción"],
    [
        ["cewb_eliminar_posicion_acero(zfer)", "Navega a CEWB → busca posición 0106/0116 → elimina la fila. Para flujo sin→sin acero"],
        ["cewb_agregar_posicion_acero(zfer)", "Navega a CEWB → agrega posición 0116 con el material de acero correspondiente. Para flujo con→con acero"],
    ],
    col_widths=[5, 13]
)
separador()

h2("Flujos de alto nivel")
tabla(
    ["Método", "Descripción"],
    [
        ["procesar_combinacion(item, res)", "Flujo completo de cambio de COLOR (7 pasos)"],
        ["procesar_formula_sin_acero(item, res)", "Flujo completo cambio fórmula sin→sin acero"],
        ["procesar_formula_con_acero(item, res)", "Flujo completo cambio fórmula con→con acero"],
        ["procesar_formula_mismo_acero(item, res)", "Flujo completo cambio fórmula sin cambio de acero"],
        ["cambiar_hoja_ruta(zfer_nuevo, id_hruta, res)", "Standalone: ca02_desasignar_hr + ca02_asignar_hr + c223 (si CA02 OK)"],
    ],
    col_widths=[5, 13]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  9. IDs SAP GUI CONFIRMADOS
# ═══════════════════════════════════════════════════════════════════════════════
h1("9. IDs SAP GUI Confirmados (sap_auto.py)")
para("Todos estos IDs fueron extraídos de grabaciones VBS del SAP GUI Recorder. "
     "Son los IDs reales de los controles SAP — si SAP cambia de versión, revisar.")
separador()

h2("Generales")
tabla(
    ["ID SAP GUI", "Descripción"],
    [
        ["wnd[0]/tbar[0]/okcd", "Campo T-Code (para /N{tcode})"],
        ["wnd[0]/sbar", "Barra de estado SAP (statusbar)"],
        ["wnd[0]/tbar[1]/btn[8]", "Ejecutar / F8 (universal en la mayoría de TCs)"],
        ["wnd[0]/tbar[0]/btn[3]", "Volver / F3"],
        ["wnd[0]/tbar[0]/btn[0]", "Guardar"],
    ],
    col_widths=[8, 10]
)
separador()

h2("ZMME0001")
tabla(
    ["ID SAP GUI", "Campo"],
    [
        ["wnd[0]/usr/radRB5", "Radio: Homologar"],
        ["wnd[0]/usr/radRB3_A1", "Radio: Cambio de Color"],
        ["wnd[0]/usr/radRB2_A1", "Radio: Cambio de Fórmula"],
        ["wnd[0]/usr/ctxtP_MATER-LOW", "Material ZFER base"],
        ["wnd[0]/usr/ctxtP_CENTER", "Centro (CO01)"],
        ["wnd[0]/usr/ctxtP_COLOR", "Código de color (ej: '19')"],
        ["wnd[0]/usr/ctxtP_FRANJ", "Código de franja (ej: '00')"],
        ["wnd[0]/usr/ctxtP_ZPLA", "ZPLA de referencia"],
        ["wnd[0]/usr/ctxtP_FORMU", "Fórmula nueva (solo cambio fórmula)"],
        ["wnd[1]/usr/cntlLO_CONTAINER0500/shellcont/shell", "Grid del popup F4 lista de ZPLAs"],
        ["wnd[0]/usr/cntlGRID1/shellcont/shell", "Grid resultado después de F8"],
        ["wnd[0]/usr/btnBUTTON1", "Botón Comparar BOM"],
        ["(TBL_BASE)/btnT_LISTA_MATERIA_INSERT", "Botón Insert en tabla inferior"],
        ["(TBL_BASE)/tblZMME0001T_LISTA_MATERIA", "Tabla inferior de posiciones BOM"],
        ["(TBL_BASE)/tblZMME0001T_LISTA_MATERIA/txtWA_LISTA-POSNR[0,{fila}]", "Columna POSNR"],
        ["(TBL_BASE)/tblZMME0001T_LISTA_MATERIA/ctxtWA_LISTA-CLASE_DESTINO[3,{fila}]", "Columna Clase Destino"],
        ["(TBL_BASE)/btnCOPY_ITEM", "Botón COPY_ITEM"],
    ],
    col_widths=[10, 8]
)
code("TBL_BASE = wnd[0]/usr/tabsTABSTRIP_MAX/tabpPUSH1/ssub%_SUBSCREEN_MAX:ZMME0001:0200")
separador()

h2("ZPPR0020")
tabla(
    ["ID SAP GUI", "Campo"],
    [
        ["wnd[0]/usr/txtS_USER-LOW", "Modificado por (PROGRAING)"],
        ["wnd[0]/usr/ctxtS_WERKS-LOW", "Centro (CO01)"],
        ["wnd[0]/usr/cntlGRID1/shellcont/shell", "Grid ALV (intenta también cntlGRID, cntlEUGRID, cntlZPPR_GRID)"],
    ],
    col_widths=[8, 10]
)
separador()

h2("MM02 — Clasificación PIEZA")
tabla(
    ["ID SAP GUI", "Campo"],
    [
        ["wnd[0]/usr/ctxtRMMG1-MATNR", "Material (número ZFER)"],
        ["wnd[0]/usr/tabsTABSPR1/tabpSP03", "Tab Clasificación"],
        ["(MM02_TAB4)", "Sub-tab PIEZA"],
        ["(MM02_TABLA)/ctxtRCTMS-MWERT[1,0]", "PARTNUMBER AGP (fila 0)"],
        ["(MM02_TABLA)/ctxtRCTMS-MWERT[1,1]", "COLOR (fila 1)"],
        ["(MM02_TABLA)/ctxtRCTMS-MWERT[1,2]", "FRANJA (fila 2)"],
        ["wnd[1]/usr/btnSPOP-OPTION1", "Confirmar diálogo de guardar (si aparece)"],
    ],
    col_widths=[10, 8]
)
code(
    "MM02_TAB4  = wnd[0]/usr/subSUBSCR_BEWERT:SAPLCTMS:5000/tabsTABSTRIP_CHAR/tabpTAB4\n"
    "MM02_TABLA = MM02_TAB4/ssubTABSTRIP_CHAR_GR:SAPLCTMS:5100/tblSAPLCTMSCHARS_S"
)
separador()

h2("CA02 — Hojas de Ruta")
tabla(
    ["ID SAP GUI", "Campo"],
    [
        ["wnd[0]/usr/ctxtRC27M-MATNR", "MATNR (material de la HR)"],
        ["wnd[0]/usr/ctxtRC27M-WERKS", "WERKS (centro CO01)"],
        ["wnd[0]/usr/ctxtRC271-PLNNR", "ID_HRUTA (número del grupo de planificación)"],
        ["wnd[0]/tbar[1]/btn[5]", "Ejecutar / Buscar"],
        ["wnd[0]/tbar[1]/btn[31]", "Abrir popup de materiales asignados"],
        ["wnd[1]/usr/tblSAPLCZDITCTRL_1010", "Tabla popup materiales"],
        ["(TBL_CA02)/txtMAPL-PLNAL[0,{row}]", "Columna PLNAL (contador)"],
        ["(TBL_CA02)/ctxtMAPL-MATNR[2,{row}]", "Columna MATNR (material asignado) ← solo funciona con findById, NO getCellValue"],
        ["(TBL_CA02)/ctxtMAPL-WERKS[3,{row}]", "Columna WERKS"],
        ["wnd[1]/tbar[1]/btn[14]", "Borrar fila (desasignación)"],
        ["wnd[1]/tbar[0]/btn[0]", "Confirmar (asignación)"],
        ["wnd[2]/tbar[0]/btn[0]", "Confirmación de borrado (nivel 2)"],
        ["wnd[2]/usr/btnSPOP-OPTION1", "Confirmación alternativa"],
        ["wnd[0]/tbar[0]/btn[11]", "Guardar"],
        ["wnd[0]/tbar[0]/btn[15]", "Volver"],
    ],
    col_widths=[9, 9]
)
code("TBL_CA02 = wnd[1]/usr/tblSAPLCZDITCTRL_1010")
separador()

h2("C223 — Versión de Fabricación")
tabla(
    ["ID SAP GUI", "Campo"],
    [
        ["wnd[0]/usr/subSUBSCR_1100:SAPLCMFV:1100/ctxtRC23M-WERKS", "Centro (CO01)"],
        ["wnd[0]/usr/subSUBSCR_1100:SAPLCMFV:1100/ctxtRC23M-MATNR", "Material ZFER"],
        ["wnd[0]/usr/ssubSUBSCR_1200:SAPLCMFV:1200/tblSAPLCMFVT_MKAL/ctxtMKAL_EXPAND-PLNNR[16,0]", "ID Hoja de Ruta (columna 16, fila 0)"],
        ["wnd[0]/usr/ssubSUBSCR_1200:SAPLCMFV:1200/tblSAPLCMFVT_MKAL/ctxtMKAL_EXPAND-MATNR[1,0]", "MATNR de la versión de fabricación"],
        ["wnd[1]/usr/btnSPOP-OPTION1", "Confirmar popup de cambio"],
        ["wnd[0]/tbar[1]/btnPRUEFEN", "Botón Verificar"],
        ["wnd[0]/tbar[1]/btn[12]", "Botón adicional (paso previo a guardar)"],
        ["wnd[0]/tbar[1]/btn[8]", "Ejecutar"],
        ["wnd[0]/tbar[0]/btn[11]", "Guardar"],
        ["wnd[0]/tbar[0]/btn[15]", "Salir"],
    ],
    col_widths=[10, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  10. PESTAÑA EXPLORADOR
# ═══════════════════════════════════════════════════════════════════════════════
h1("10. Pestaña Explorador (/explorar y /zfer/<material>)")
para(
    "El explorador es la herramienta de consulta del sistema. Permite buscar cualquier ZFER "
    "por múltiples criterios y ver su ficha completa con todos los atributos."
)
separador()

h2("Búsqueda (/explorar)")
para("Acepta estos parámetros vía GET (todos opcionales, se combinan con AND):")
tabla(
    ["Parámetro URL", "Campo SAP", "Tipo de búsqueda"],
    [
        ["vehiculo", "Z_VEHICLE_MODEL", "LIKE %valor%"],
        ["formula", "Z_FORMULA_CODE", "LIKE %valor%"],
        ["pieza", "Z_PIECE_TYPE", "LIKE %valor%"],
        ["color", "Z_COLOR", "LIKE %valor%"],
        ["version", "Z_AGP_VERSION", "LIKE %valor%"],
        ["nivel", "Z_AGP_LEVEL", "LIKE %valor%"],
        ["cod_vehiculo", "Z_AGP_PARTNUMBER", "LIKE {cod_vehiculo}_%  (prefijo del PN)"],
    ],
    col_widths=[4, 5, 9]
)
para("Máximo 300 resultados. La búsqueda usa un OR de ATNAM=? AND ATWRT LIKE ? con GROUP BY/HAVING "
     "para intersectar los criterios en un solo scan.", italic=True)
separador()

h2("Ficha de ZFER (/zfer/<material>)")
para("Muestra toda la información disponible del ZFER. Datos cargados en paralelo con ThreadPoolExecutor:")
tabla(
    ["Sección", "Fuente de datos", "Función"],
    [
        ["Atributos básicos", "ODATA_ZFER_HEAD", "q_zfer_head()"],
        ["Clasificación SAP", "ODATA_ZFER_CLASS_001", "q_atributos()"],
        ["Variantes de color", "ODATA_ZFER_CLASS_001 + HEAD (JOIN)", "q_variantes_por_pn()"],
        ["Entregas y mercados", "ODATA_ZCDS_Entregas_Pos_CO + Head_CO", "q_entregas() + q_mercados()"],
        ["ZPLAs compatibles", "ODATA_ZPLA_CLASS_001 + HEAD", "q_zplas_compatibles()"],
        ["Fórmulas alternativas", "ODATA_ZPLA_HEAD + CLASS_001 (CTE)", "q_formulas_por_pieza()"],
        ["Plano técnico", "ODATA_ZFER_RUTAS_JPG", "_q_plano() + api_plano()"],
        ["Par simétrico", "Calculado desde _PARES_SIMETRIA", "_PARES_SIMETRIA dict"],
    ],
    col_widths=[4, 6, 5, 3]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  11. PESTAÑA HOJAS DE RUTA
# ═══════════════════════════════════════════════════════════════════════════════
h1("11. Pestaña Hojas de Ruta (/hojas_ruta)")
para(
    "Esta página busca la hoja de ruta de producción más adecuada para un ZFER nuevo "
    "y permite asignarla directamente en SAP con un click."
)
separador()

h2("Flujo de la página")
bullet("Usuario ingresa ZFER base + ZFER nuevo")
bullet("Sistema consulta BD para ZFER base: NIVEL, GEOMETRIA (Z_GEOMETRY_TYPE), AREA (→ TAMAÑO)")
bullet("Sistema consulta SAP (ZPPR0008) para ZFER nuevo: posiciones del BOM")
bullet("Con esos datos construye criterios → filtra ODATA_HR_CONSULTA")
bullet("Muestra candidata (MAX materiales ≤ 450) y todas las demás opciones")
separador()

h2("Criterios de filtrado (_hr_construir_criterios)")
tabla(
    ["Criterio", "Fuente", "Lógica"],
    [
        ["NIVEL", "Z_AGP_LEVEL del ZFER base", "≤3 → BAJO, 4 → MEDIO, ≥5 → ALTO"],
        ["GEOMETRIA", "Z_GEOMETRY_TYPE del ZFER base", "02 → CURVO, 01 → PLANO"],
        ["TAMAÑO", "AREA del ZFER base (ODATA_ZFER_HEAD)", "≤0.6 → PEQUEÑA, ≤0.99 → MEDIANA, ≥1.00 → GRANDE"],
        ["FORMULA", "Z_FORMULA_CODE del ZFER base", "Mapa _HR_CLAVE_FORMULA: posiciones 100-800"],
        ["BASE", "Posición 99 en BOM del ZFER nuevo", "32VPMO si está presente"],
        ["PROTECTORS", "Posiciones 199/299 en BOM", "33VPR01 / 34VPR02"],
        ["TAPAS", "Posiciones 3600/3700 en BOM", "36VTPA / 37VSTP"],
        ["EMPALME", "CURVO: suma formula+base+protectors+tapas", "Si BOM vacío y CURVO → IS NOT NULL"],
        ["CURVADO", "Si CURVO=True", "1 si curvo, NULL si plano"],
        ["CURV_ACERO", "Pos 106 o 116 en BOM", "1 si tiene acero, NULL si no"],
        ["METROLOGIA", "HR del ZFER base", "Heredado"],
        ["PRUEBA_AGUA", "HR del ZFER base", "Heredado"],
        ["SERIGRAFIA", "Criterio del ZFER", "excluir_null=True (no trae HRs con NULL cuando el criterio tiene valor)"],
        ["MECANIZADO", "Criterio del ZFER", "excluir_null=True"],
        ["VITRIFICADO", "Criterio del ZFER", "excluir_null=True"],
    ],
    col_widths=[3, 5, 10]
)
separador()

h2("Selección de candidata")
para("La candidata es la HR con MAYOR cantidad de materiales asignados entre las que cumplen "
     "todos los criterios, siempre que TOTAL_MATERIALES ≤ 450.")
para("El límite de 450 garantiza que la HR no esté saturada. La barra de capacidad en la UI "
     "va de verde (0) a roja (cerca de 450).", italic=True)
separador()

h2("Asignación (botón Asignar HR)")
bullet("Llama ca02_desasignar_hr(zfer_nuevo) — quita la HR actual si tiene una")
bullet("Llama ca02_asignar_hr(zfer_nuevo, id_hruta) — asigna la nueva")
bullet("Si CA02 retorna True → llama c223_actualizar_version_fabricacion(zfer_nuevo, id_hruta)")
bullet("Guarda el resultado en dbo.M5_RutasZFER")
separador()

h2("Validador de antenas pasta plata")
para("Las posiciones 9452-9456 en el BOM son de pasta plata (antenas). "
     "El sistema verifica que el último dígito del número de posición coincida con "
     "el primer dígito de la última posición de fórmula del BOM. "
     "Si no coincide → alerta amarilla en la UI (no bloquea el proceso).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  12. PESTAÑA COLA
# ═══════════════════════════════════════════════════════════════════════════════
h1("12. Pestaña Cola de Homologaciones (/cola)")
para(
    "La cola es el centro de operaciones SAP. Acá se organizan los trabajos, "
    "se programan para ejecutarse a una hora específica y se monitorea el progreso en tiempo real."
)
separador()

h2("Estructura de un Bloque")
para("Un BLOQUE es un contenedor de tareas. Puede tener N items (uno por ZFER a procesar). "
     "Estados posibles de un bloque:")
tabla(
    ["Estado", "Descripción"],
    [
        ["PENDIENTE", "Creado pero no ejecutado todavía. Puede tener timer programado."],
        ["EJECUTANDO", "Actualmente procesando items en SAP."],
        ["COMPLETADO", "Todos los items terminaron (con o sin errores individuales)."],
        ["ERROR", "El bloque falló de forma catastrófica (ej: SAP no disponible)."],
    ],
    col_widths=[4, 14]
)
separador()

h2("Scheduler automático")
para("El scheduler corre como hilo daemon desde que arranca Flask. "
     "Cada 20 segundos revisa si hay bloques con timer_activo=1 cuya hora_prog ya llegó.")
code(
    "_cola_scheduler():\n"
    "  _cola_limpiar_al_inicio()   # recovery: EJECUTANDO → PENDIENTE al arrancar\n"
    "  _cola_scheduler_tick()      # tick inmediato al iniciar\n"
    "  while True:\n"
    "      sleep(20)\n"
    "      _cola_scheduler_tick()\n\n"
    "_cola_scheduler_tick():\n"
    "  # Consulta M5_Bloques WHERE estado='PENDIENTE' AND timer_activo=1 AND hora_prog IS NOT NULL\n"
    "  # Para cada bloque vencido: si NO está en _scheduler_disparados → disparar en hilo nuevo\n"
    "  # _scheduler_disparados: set() para dedup (evita doble disparo si ciclos se solapan)"
)
separador()

h2("Tipos de item en la cola (campo 'tipo' en M5_Cola)")
tabla(
    ["Tipo", "Flujo SAP ejecutado"],
    [
        ["color", "procesar_combinacion() — Cambio de color"],
        ["FORMULA_SIN_ACERO", "procesar_formula_sin_acero() — Cambio fórmula sin acero → sin acero"],
        ["FORMULA_CON_ACERO", "procesar_formula_con_acero() — Cambio fórmula con acero → con acero"],
        ["FORMULA_MISMO_ACERO", "procesar_formula_mismo_acero() — Cambio fórmula sin cambio de acero"],
    ],
    col_widths=[5, 13]
)
separador()

h2("Reporte Excel (descarga)")
para("Al completar un bloque se puede descargar un Excel con 4 hojas:")
tabla(
    ["Hoja", "Contenido"],
    [
        ["RESUMEN", "Info del bloque: total items, OK, errores, tiempo, usuario"],
        ["PROCESADOS_SAP", "Todos los items con ZFER nuevo, estado, advertencias (OK=verde, ERROR=rojo)"],
        ["SOLO_REPORTE", "Items que no fueron a SAP (fórmula diferente) — en amarillo"],
        ["ERRORES", "Solo los errores para revisión rápida del técnico SAP"],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  13. PESTAÑA COMBINACIONES
# ═══════════════════════════════════════════════════════════════════════════════
h1("13. Pestaña Combinaciones (/combinaciones)")
para(
    "Genera todas las combinaciones posibles de color para un ZFER base, "
    "mostrando qué ZPLAs están disponibles para cada color."
)
separador()
h2("Cómo funciona")
bullet("Usuario ingresa un ZFER base")
bullet("Sistema llama q_atributos() para obtener: Z_FORMULA_CODE, Z_PIECE_TYPE, Z_SHADE_BAND, Z_BEHAVIOR_DIFFERENTIALS")
bullet("Con esos datos llama q_zplas_compatibles() → lista de ZPLAs × colores disponibles")
bullet("Filtra por _COLORES_ACTIVOS (solo los 12 colores habilitados)")
bullet("Muestra en una cuadrícula de 4 columnas: Incoloro / Claros / Medio / Dark")
bullet("Usuario selecciona combinaciones → 'Enviar a Cola' → se crea un bloque en M5_Bloques con los items seleccionados")
separador()
h2("Columnas de la cuadrícula de colores")
tabla(
    ["Columna", "Colores incluidos"],
    [
        ["Incoloro", "00 (Blanco)"],
        ["Claros", "01, 05, 06"],
        ["Medio", "10, 13, 18, 19, 20, 22"],
        ["Dark", "21, 23"],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  14. TIMINGS Y RENDIMIENTO
# ═══════════════════════════════════════════════════════════════════════════════
h1("14. Timings y Rendimiento")
para(
    "El sistema fue optimizado para ser lo más rápido posible sin romper la lógica SAP. "
    "Los tiempos son máximos — la espera real termina en cuanto SAP responde (app.Busy = False)."
)
separador()

h2("Constantes de tiempo (sap_auto.py)")
tabla(
    ["Constante", "Valor", "Uso"],
    [
        ["T_RAPIDO", "0.5 seg", "Máximo para clicks simples y campos de texto"],
        ["T_MEDIO", "1.5 seg", "Máximo para navegación entre pantallas"],
        ["T_LENTO", "6.0 seg", "Máximo para ejecutar transacciones pesadas (F8, F4)"],
        ["_T_MIN_RAPIDO", "0.02 seg", "Mínimo garantizado antes de empezar el poll (T_RAPIDO)"],
        ["_T_MIN_MEDIO", "0.03 seg", "Mínimo garantizado (T_MEDIO)"],
        ["_T_MIN_LENTO", "0.05 seg", "Mínimo garantizado (T_LENTO)"],
        ["_T_POLL", "0.05 seg", "Intervalo de poll de app.Busy"],
    ],
    col_widths=[4, 3, 11]
)
separador()

h2("Optimizaciones especiales")
tabla(
    ["Área", "Optimización", "Impacto"],
    [
        ["CA02 scroll", "time.sleep(0.02) en lugar de _esperar(T_RAPIDO). Block-skip: sp += max(1, vis_rows-1)", "1,487 filas: de ~11,000 a ~186 findById calls"],
        ["Diferencial popup", "Paginación por bloques (vis_pop=10 → 5 scrolls para 43 items). Sin setFocus.", "Antes: 43 scrolls individuales + event race conditions"],
        ["Planos bulk", "_q_planos_bulk(): un solo IN query para N materiales", "De N queries individuales a 1 solo"],
        ["ZPPR0020 fases", "60 intentos máx (antes 120), 5s entre intentos", "Tiempo máximo: 300s en lugar de 600s"],
        ["Connection pool", "_conn_pool Queue(maxsize=12) con health-check SELECT 1", "Elimina reconexiones TCP por cada query"],
        ["Atributos paralelos", "ThreadPoolExecutor en /zfer/<material>: carga q_zfer_head, q_atributos, q_variantes_por_pn, q_entregas en paralelo", "Reducción del tiempo de carga de la ficha ZFER"],
        ["Scheduler", "Tick cada 20s + tick inmediato al arrancar + dedup con set()", "No pierde alarmas, no duplica disparos"],
    ],
    col_widths=[3, 7, 8]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  15. USUARIOS Y SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════
h1("15. Usuarios y Seguridad")
para(
    "El sistema usa autenticación simple por email + contraseña almacenada en un dict Python "
    "(no en BD). La sesión se maneja con Flask session (cookie firmada con secret_key)."
)
separador()
h2("Usuarios configurados")
para("Los usuarios están definidos en _USUARIOS dict en app.py. Algunos destacados:")
tabla(
    ["Email", "Rol"],
    [
        ["atcol@agpglass.com", "Admin / IT Colombia"],
        ["fguerrero@agpglass.com", "Ingeniería"],
        ["spina@agpglass.com", "Ingeniería de Producto"],
        ["pract1-4@agpglass.com", "Practicantes de Ingeniería"],
        ["leo@agpglass.com / prueba@agpglass.com", "Cuentas de prueba"],
    ],
    col_widths=[7, 11]
)
separador()
h2("Decorator login_required")
para("Todas las rutas excepto /login y /logout requieren sesión activa. "
     "Si no hay sesión → redirect a /login con el parámetro next= para volver después.")
para("Secret key Flask: AGP_M5_2025_xK9!mQ#zL (firmado criptográficamente en la cookie).", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  16. ESTRUCTURA ResultadoItem
# ═══════════════════════════════════════════════════════════════════════════════
h1("16. Estructura ResultadoItem (sap_auto.py)")
para(
    "Es el objeto que viaja con cada tarea SAP durante toda su ejecución. "
    "Acumula el log, errores, advertencias y el resultado final."
)
separador()
tabla(
    ["Campo", "Tipo", "Descripción"],
    [
        ["batch_id", "str", "UUID único del lote de procesamiento"],
        ["zfer_base", "str", "ZFER de entrada"],
        ["color_codigo", "str", "Código del color SAP"],
        ["zfer_nuevo", "str", "ZFER creado por SAP (se llena en PASO 2)"],
        ["zfor_nuevo", "str", "ZFOR creado (si aplica)"],
        ["zpla", "str", "ZPLA seleccionado en el proceso"],
        ["posiciones_bom", "list", "Lista de posiciones BOM procesadas (ej: ['0458'])"],
        ["bom_detalle", "list", "Lista de {posnr, clase_destino} para cada posición"],
        ["advertencias", "list", "Advertencias no-fatales: plano no encontrado, HR sin candidata, etc."],
        ["estado", "str", "PENDIENTE / EN_PROCESO / OK / ERROR"],
        ["error", "str", "Mensaje de error si estado=ERROR"],
        ["fecha_inicio", "datetime", "Timestamp inicio del procesamiento"],
        ["fecha_fin", "datetime", "Timestamp fin del procesamiento"],
        ["log", "list", "Log detallado de todos los pasos (para debugging)"],
        ["formula", "str", "Fórmula del ZFER base (para log BD)"],
        ["tipo_pieza", "str", "Tipo de pieza (para log BD)"],
        ["acero", "str", "Descripción del acero (para log BD)"],
        ["color_nombre", "str", "Nombre del color (para log BD)"],
        ["tipo", "str", "color / formula (para log BD)"],
        ["duracion_seg", "float (property)", "Calculado: (fecha_fin - fecha_inicio).total_seconds()"],
    ],
    col_widths=[4, 3, 11]
)
separador()
h2("Métodos de ResultadoItem")
tabla(
    ["Método", "Descripción"],
    [
        ["_log(msg)", "Imprime en consola con prefijo [SAP] y agrega al list log[]"],
        ["_advertir(msg)", "Imprime con prefijo [SAP][ADV], agrega a advertencias[] con prefijo [ADV] y también al log[]. Los items en advertencias[] se muestran en amarillo en el reporte Excel."],
    ],
    col_widths=[4, 14]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
#  APÉNDICE — DEPENDENCIAS RÁPIDAS
# ═══════════════════════════════════════════════════════════════════════════════
h1("Apéndice — Mapa de Dependencias Rápido")
para("Tabla de referencia: qué función lee qué tabla de BD.")
separador()
tabla(
    ["Función / Método", "Lee de BD", "Escribe en BD"],
    [
        ["q_zfer_head()", "ODATA_ZFER_HEAD", "—"],
        ["q_atributos()", "ODATA_ZFER_CLASS_001", "—"],
        ["q_entregas()", "ODATA_ZCDS_Entregas_Pos_CO", "—"],
        ["q_mercados()", "ODATA_ZCDS_Entregas_Head_CO", "—"],
        ["q_variantes_por_pn()", "ODATA_ZFER_CLASS_001 + ODATA_ZFER_HEAD", "—"],
        ["q_zplas_compatibles()", "ODATA_ZPLA_CLASS_001 + ODATA_ZPLA_HEAD", "—"],
        ["q_formulas_por_pieza()", "ODATA_ZPLA_HEAD + ODATA_ZPLA_CLASS_001", "—"],
        ["q_explorar()", "ODATA_ZFER_CLASS_001 + ODATA_ZFER_HEAD", "—"],
        ["_q_plano() / _q_planos_bulk()", "ODATA_ZFER_RUTAS_JPG", "—"],
        ["_hr_buscar_candidata/todas()", "ODATA_HR_CONSULTA", "—"],
        ["_hr_construir_criterios()", "ODATA_ZFER_HEAD, ODATA_ZFER_CLASS_001, ODATA_HR_CONSULTA", "—"],
        ["_cola_ejecutar_bloque()", "M5_Cola, M5_Bloques", "M5_Cola (estado), M5_Bloques (estado, contadores), M5_LogEjecucion"],
        ["_cola_scheduler_tick()", "M5_Bloques", "—"],
        ["_migracion_bd_local()", "—", "M5_Cola, M5_Bloques (ALTER TABLE ADD si falta columna)"],
        ["mm02_actualizar_diferenciales_zpla()", "ODATA_ZPLA_CLASS_001 (diferencial ZPLA base + orden completo)", "—"],
        ["mm02_cambiar_plano()", "ODATA_ZFER_RUTAS_JPG", "—"],
        ["_buscar_diferenciales_zpla()", "ODATA_ZPLA_CLASS_001 (BD producción agpcol)", "—"],
    ],
    col_widths=[5, 8, 5]
)

separador()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_final = p_final.add_run(f"Documento generado el {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')} — AGP Glass Colombia — Módulo 5")
run_final.italic = True
run_final.font.color.rgb = RGBColor.from_string("888888")

# ── Guardar ───────────────────────────────────────────────────────────────────
import os
output_path = os.path.join(
    r"C:\Users\abotero\OneDrive - AGP GROUP\Documentos\MODULO_5",
    "MODULO5_Documentacion_Completa.docx"
)
doc.save(output_path)
print(f"\nDocumento guardado en: {output_path}")
